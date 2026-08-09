"""DOCX 文档解析器 — 使用 python-docx 提取正文和表格内容。

处理流程：
  1. 用 python-docx 打开 .docx 文件
  2. 提取所有非空段落文本
  3. 遍历所有表格，将每行转为管道符分隔的文本（保留表格结构）
  4. 将段落文本和表格文本合并
  5. 使用 RecursiveCharacterTextSplitter 分块
  6. 从 Heading 样式提取标题树，写回 ParseResult.heading_tree

表格提取策略：
  金融文档（年报、审计报告）中大量数据以表格形式存在，
  本解析器将表格转为完整的 Markdown 表格（含表头和 |---| 分隔行），
  保留表格的结构化信息，便于 LLM 理解数字数据。

延迟导入说明：
  python-docx 在 parse() 内部才 import，避免模块顶层依赖，
  这样在未安装 python-docx 时其他解析器仍可正常使用。
"""

import os
import re

from langchain_text_splitters import RecursiveCharacterTextSplitter

from src.config import CHUNK_OVERLAP, CHUNK_SIZE
from src.parsers.base import BaseParser, ChunkData, ParseResult

# 匹配 Markdown 表格（以 | 开头和结尾的行组成的多行表格）
TABLE_PATTERN = re.compile(r"^\|.+\|[\s\S]*?^\|.+\|", re.MULTILINE)


class DocxParser(BaseParser):
    """DOCX 文档解析器 — 提取段落文本 + 表格数据。"""

    def parse(self, file_path: str) -> ParseResult:
        """解析 DOCX 文件并分块。

        Args:
            file_path: DOCX 文件路径

        Returns:
            ParseResult，file_type="docx"，total_pages=1，
            heading_tree=标题层级列表

        Raises:
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # 延迟导入 python-docx（避免模块级依赖）
        from docx import Document

        doc = Document(file_path)

        # ====== 提取段落文本 ======
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        # ====== 提取表格数据（Markdown 格式）======
        table_texts = self._extract_tables(doc)

        # 将表格文本追加到正文之后（用双换行分隔）
        if table_texts:
            text += "\n\n" + "\n\n".join(table_texts)

        # ====== 提取标题树（Heading 样式）======
        heading_tree = self._extract_heading_tree(doc)

        source = os.path.basename(file_path)
        # 分块：与 TXT 使用相同的中文友好分隔符
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", "；", " ", ""],
        )
        texts = splitter.split_text(text)
        chunks = []
        for i, t in enumerate(texts):
            block_type = "table" if TABLE_PATTERN.search(t) else "text"
            chunks.append(
                ChunkData(
                    content=t,
                    metadata={"source": source, "page": 1, "block_type": block_type},
                    chunk_id=f"{source}:{i}",
                )
            )
        return ParseResult(
            chunks=chunks,
            total_pages=1,
            total_chars=len(text),
            file_type="docx",
            heading_tree=heading_tree,
        )

    def _extract_heading_tree(self, doc) -> list[tuple[int, str]]:
        """从段落样式提取标题树（(层级, 标题) 列表）。

        识别 python-docx 的 Heading 系列样式：
          - 英文样式名: "Heading 1" / "Heading 2" ...
          - 中文样式名: "标题 1" / "标题 2" ...
        仅统计非空段落，正文段落（Normal）不进入标题树。

        Args:
            doc: python-docx Document 对象

        Returns:
            标题层级列表，元素为 (level, heading)
        """
        heading_tree: list[tuple[int, str]] = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style = para.style
            if style is not None:
                style_name = style.name or ""
            else:
                style_name = ""
            # 英文样式: "Heading 1"，取样式名最后一段作为层级数字
            if style_name.lower().startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1
                heading_tree.append((level, para.text.strip()))
            # 中文样式: "标题 1"，用正则提取层级数字
            elif "标题" in style_name:
                m = re.search(r"(\d)", style_name)
                if m:
                    level = int(m.group(1))
                else:
                    level = 1
                heading_tree.append((level, para.text.strip()))
        return heading_tree

    def _extract_tables(self, doc) -> list[str]:
        """提取 DOCX 文档中的所有表格，返回 Markdown 格式字符串列表。

        Args:
            doc: python-docx Document 对象

        Returns:
            Markdown 表格字符串列表（每个元素是一个完整的 Markdown 表格）
        """

        result = []
        for table in doc.tables:
            md = self._docx_table_to_markdown(table)
            if md:
                result.append(md)
        return result

    def _docx_table_to_markdown(self, table) -> str:
        """将 python-docx Table 对象转换为 Markdown 格式字符串。

        Args:
            table: python-docx Table 对象

        Returns:
            Markdown 格式的表格字符串，空表格返回空字符串
        """
        rows = []
        for row in table.rows:
            cells = [self.sanitize_cell(cell.text) for cell in row.cells]
            rows.append(cells)
        if not rows:
            return ""
        lines = []
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
